#!/usr/bin/env python3
# main.py

import os
import re
from io import BytesIO
from typing import List, Optional
from datetime import datetime, timedelta

from fastapi import (
    FastAPI, UploadFile, File, Form,
    Depends, HTTPException, status
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordRequestForm
from pydantic import BaseModel
from sqlalchemy.orm import Session
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import RGBColor
from openai import OpenAI

from db import init_db, get_db
from auth import (
    UserCreate, create_access_token,
    verify_password, get_current_user,
    ACCESS_TOKEN_EXPIRE_MINUTES, hash_password
)
from models import (
    User, SubscriptionPlan,
    UserSubscription, EvaluationRun
)
from resume_evaluator_api import (
    load_api_key, evaluate_resume,
    generate_interview_questions
)

app = FastAPI()

# ─── Startup & Health ────────────────────────────────────────────────────────

@app.on_event("startup")
def on_startup():
    init_db()

@app.get("/health")
async def health():
    return {"status": "ok"}


# ─── CORS ────────────────────────────────────────────────────────────────────

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://resumeyval.com",
        "https://www.resumeyval.com"
    ],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ─── Patterns & OpenAI Client ────────────────────────────────────────────────

ZERO_WIDTH = re.compile(r'[\u200B\u200C\u200D\uFEFF]')
WATERMARK  = re.compile(r'\bCONFIDENTIAL\b', re.IGNORECASE)
WHITE_PDF  = re.compile(r'1\s+1\s+1\s+(rg|RG)')

API_KEY = load_api_key()
client  = OpenAI(api_key=API_KEY)


# ─── Pydantic Schemas ────────────────────────────────────────────────────────

class PlanOut(BaseModel):
    id: int
    name: str
    description: str
    price_usd: float
    max_runs: Optional[int]

    class Config:
        orm_mode = True

class SubscriptionOut(BaseModel):
    plan: PlanOut
    starts_at: datetime
    expires_at: Optional[datetime]

    class Config:
        orm_mode = True

class MeOut(BaseModel):
    email: str
    subscription: SubscriptionOut

    class Config:
        orm_mode = True

class SubscribeIn(BaseModel):
    plan_name: str


# ─── AUTH ROUTES ─────────────────────────────────────────────────────────────

@app.post("/signup", status_code=201)
def signup(
    user_in: UserCreate,
    db: Session = Depends(get_db)
):
    # ensure email isn’t taken
    if db.query(User).filter_by(email=user_in.email).first():
        raise HTTPException(400, "Email already registered")

    # create user
    user = User(
        email     = user_in.email,
        hashed_pw = hash_password(user_in.password)
    )
    db.add(user)
    db.commit()
    db.refresh(user)

    # assign free plan
    free_plan = db.query(SubscriptionPlan)\
                  .filter_by(name="free").one()
    sub = UserSubscription(
        user_id = user.id,
        plan_id = free_plan.id,
        starts_at = datetime.utcnow(),
        expires_at = None
    )
    db.add(sub)
    db.commit()

    return {"msg": "User created and assigned free plan"}


@app.post("/token")
def login_for_access_token(
    form_data: OAuth2PasswordRequestForm = Depends(),
    db: Session = Depends(get_db)
):
    user = db.query(User)\
             .filter_by(email=form_data.username)\
             .first()
    if not user or not verify_password(
        form_data.password, user.hashed_pw
    ):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )

    token = create_access_token(
        data={ "sub": user.email },
        expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    )
    return { "access_token": token, "token_type": "bearer" }


# ─── SUBSCRIPTION MANAGEMENT ─────────────────────────────────────────────────

@app.get("/plans", response_model=List[PlanOut])
def list_plans(db: Session = Depends(get_db)):
    """List all available subscription plans."""
    return db.query(SubscriptionPlan).all()


@app.get("/me", response_model=MeOut)
def read_current_user(current_user: User = Depends(get_current_user)):
    """Get current user’s profile and subscription."""
    return current_user


@app.post("/subscribe", response_model=SubscriptionOut)
def subscribe(
    req: SubscribeIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Switch the current user’s plan."""
    plan = db.query(SubscriptionPlan)\
             .filter_by(name=req.plan_name)\
             .first()
    if not plan:
        raise HTTPException(404, "Plan not found")

    # update subscription record
    sub: UserSubscription = current_user.subscription
    sub.plan_id    = plan.id
    sub.starts_at  = datetime.utcnow()
    # free or perpetual plans never expire
    sub.expires_at = None

    db.add(sub)
    db.commit()
    db.refresh(sub)
    return sub


# ─── EVALUATION (protected + usage limits) ─────────────────────────────────

@app.post("/evaluate/")
async def evaluate(
    job_description: str = Form(...),
    resumes: List[UploadFile] = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    plan = current_user.subscription.plan

    # enforce max_runs limit
    if plan.max_runs is not None:
        used = db.query(EvaluationRun)\
                 .filter_by(user_id=current_user.id)\
                 .count()
        if used >= plan.max_runs:
            raise HTTPException(
                403,
                detail="Run limit reached; please upgrade your plan."
            )

    results = []
    for up in resumes:
        fname = up.filename
        ext   = os.path.splitext(fname)[1].lower()
        raw   = await up.read()
        flags = []

        # 1) extract + detect tricks
        try:
            if ext == ".pdf":
                reader = PdfReader(BytesIO(raw))
                text   = "\n".join(
                    p.extract_text() or "" for p in reader.pages
                )
                if WHITE_PDF.search(raw.decode("latin-1","ignore")):
                    flags.append("Hidden white text in PDF")

            elif ext == ".docx":
                doc = Document(BytesIO(raw))
                for para in doc.paragraphs:
                    for run in para.runs:
                        c = run.font.color
                        if c and c.rgb == RGBColor(0xFF,0xFF,0xFF):
                            flags.append("Hidden white text in DOCX")
                            break
                    if "Hidden white text in DOCX" in flags:
                        break
                text = "\n".join(p.text for p in doc.paragraphs)

            elif ext == ".txt":
                text = raw.decode("utf-8","ignore")

            else:
                raise ValueError(f"Unsupported: {ext}")

            if ZERO_WIDTH.search(text):
                flags.append("Hidden/invisible characters")
            if WATERMARK.search(text):
                flags.append("Watermark 'CONFIDENTIAL' detected")

        except Exception as e:
            results.append({
                "filename": fname,
                "error": f"Extraction failed: {e}"
            })
            continue

        # 2) call OpenAI
        try:
            eval_data  = evaluate_resume(
                client, "gpt-4", job_description, text
            )
            questions = generate_interview_questions(
                client, "gpt-4", job_description, text
            )

            summary = (
                f"{eval_data['name']} - "
                f"{eval_data['match_score']}% match - "
                f"{eval_data['verdict']} "
                f"[{'green flag' if eval_data['verdict']=='Strong Fit' else 'red flag'}]"
            )
            reasons = (eval_data.get("green_flags",[])
                       + eval_data.get("red_flags",[]))[:4]

            # free plan: minimal detail
            if plan.name == "free":
                reasons   = reasons[:2]
                questions = []

            results.append({
                "filename": fname,
                "summary": summary,
                "reasons": reasons,
                "interview_questions": questions[:5],
                "suspicious_flags": flags
            })

        except Exception as e:
            results.append({
                "filename": fname,
                "error": str(e)
            })

    # record that we used one run
    run = EvaluationRun(user_id=current_user.id)
    db.add(run)
    db.commit()

    return {"results": results}
